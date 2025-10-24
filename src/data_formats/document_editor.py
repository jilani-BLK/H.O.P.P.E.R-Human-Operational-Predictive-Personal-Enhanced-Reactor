"""
Éditeur de documents pour formats complexes
Modification sécurisée de PDF, DOCX, Excel, CSV sans corruption
"""

from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field
from pathlib import Path
from enum import Enum
import json
import csv
import shutil
from datetime import datetime

# Import from readers
try:
    from PyPDF2 import PdfReader, PdfWriter
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    from openpyxl import load_workbook
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False


class EditOperationType(Enum):
    """Types d'opérations d'édition"""
    INSERT = "insert"
    DELETE = "delete"
    REPLACE = "replace"
    APPEND = "append"
    UPDATE_CELL = "update_cell"  # Pour tableaux
    ADD_PAGE = "add_page"  # Pour PDF/DOCX
    REMOVE_PAGE = "remove_page"
    MERGE = "merge"  # Fusionner documents


@dataclass
class EditOperation:
    """Opération d'édition"""
    operation_type: EditOperationType
    target: str  # Localisation: "page:2", "paragraph:5", "cell:A1", etc.
    content: Optional[Any] = None
    options: Dict[str, Any] = field(default_factory=dict)


@dataclass
class EditResult:
    """Résultat d'édition"""
    success: bool
    operations_applied: int = 0
    backup_path: Optional[str] = None
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    modified_sections: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return {
            "success": self.success,
            "operations_applied": self.operations_applied,
            "backup_path": self.backup_path,
            "warnings": self.warnings,
            "errors": self.errors,
            "modified_sections": self.modified_sections
        }


class DocumentEditor:
    """Éditeur sécurisé de documents"""
    
    def __init__(self, llm_analyzer: Optional[Any] = None):
        self.llm = llm_analyzer
        self.backup_enabled = True
        self.backup_dir = Path("data/backups")
        self.backup_dir.mkdir(parents=True, exist_ok=True)
    
    async def edit_document(
        self,
        file_path: str,
        operations: List[EditOperation],
        create_backup: bool = True
    ) -> EditResult:
        """Applique des modifications à un document"""
        
        source = Path(file_path)
        if not source.exists():
            return EditResult(
                success=False,
                errors=["Fichier introuvable"]
            )
        
        # Backup automatique
        backup_path = None
        if create_backup and self.backup_enabled:
            backup_path = self._create_backup(source)
        
        # Détecte format
        file_format = self._detect_format(source)
        
        # Route vers éditeur approprié
        try:
            if file_format == "pdf":
                result = await self._edit_pdf(source, operations)
            elif file_format == "docx":
                result = await self._edit_docx(source, operations)
            elif file_format == "xlsx":
                result = await self._edit_excel(source, operations)
            elif file_format == "csv":
                result = await self._edit_csv(source, operations)
            elif file_format in ["txt", "md", "json"]:
                result = await self._edit_text_file(source, operations)
            else:
                result = EditResult(
                    success=False,
                    errors=[f"Format {file_format} non supporté pour édition"]
                )
            
            result.backup_path = backup_path
            return result
            
        except Exception as e:
            # Restaure backup en cas d'erreur
            if backup_path:
                shutil.copy2(backup_path, source)
            
            return EditResult(
                success=False,
                backup_path=backup_path,
                errors=[f"Erreur édition: {str(e)}"]
            )
    
    def _create_backup(self, file_path: Path) -> str:
        """Crée une sauvegarde du fichier"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"{file_path.stem}_{timestamp}{file_path.suffix}"
        backup_path = self.backup_dir / backup_name
        
        shutil.copy2(file_path, backup_path)
        return str(backup_path)
    
    def _detect_format(self, file_path: Path) -> str:
        """Détecte le format du fichier"""
        suffix = file_path.suffix.lower().lstrip('.')
        format_map = {
            'pdf': 'pdf',
            'docx': 'docx',
            'xlsx': 'xlsx',
            'csv': 'csv',
            'json': 'json',
            'md': 'md',
            'txt': 'txt'
        }
        return format_map.get(suffix, 'unknown')
    
    async def _edit_pdf(
        self,
        file_path: Path,
        operations: List[EditOperation]
    ) -> EditResult:
        """Édite un PDF"""
        if not HAS_PDF:
            return EditResult(
                success=False,
                errors=["PyPDF2 non installé"]
            )
        
        try:
            reader = PdfReader(str(file_path))
            writer = PdfWriter()
            
            modified_sections = []
            operations_applied = 0
            warnings = []
            
            # Copie pages existantes
            pages_to_keep = list(range(len(reader.pages)))
            
            for op in operations:
                if op.operation_type == EditOperationType.REMOVE_PAGE:
                    # Format: "page:2" (commence à 1)
                    page_num = int(op.target.split(':')[1]) - 1
                    if 0 <= page_num < len(reader.pages):
                        pages_to_keep.remove(page_num)
                        modified_sections.append(f"Supprimé page {page_num + 1}")
                        operations_applied += 1
                    else:
                        warnings.append(f"Page {page_num + 1} n'existe pas")
                
                elif op.operation_type == EditOperationType.ADD_PAGE:
                    # Ajoute page vide ou depuis autre PDF
                    if isinstance(op.content, str) and Path(op.content).exists():
                        # Fusionne depuis autre PDF
                        source_pdf = PdfReader(op.content)
                        for page in source_pdf.pages:
                            writer.add_page(page)
                        modified_sections.append(f"Ajouté pages depuis {op.content}")
                        operations_applied += 1
                    else:
                        warnings.append("Ajout de page vide non implémenté")
            
            # Ajoute pages conservées
            for page_num in pages_to_keep:
                writer.add_page(reader.pages[page_num])
            
            # Note: Modification de contenu texte PDF complexe
            # Nécessite bibliothèque avancée (pikepdf, PyMuPDF)
            for op in operations:
                if op.operation_type in [EditOperationType.INSERT, EditOperationType.REPLACE]:
                    warnings.append(
                        "Modification de texte PDF nécessite pikepdf/PyMuPDF "
                        "(non inclus par défaut)"
                    )
            
            # Sauvegarde
            with open(file_path, 'wb') as f:
                writer.write(f)
            
            return EditResult(
                success=True,
                operations_applied=operations_applied,
                warnings=warnings,
                modified_sections=modified_sections
            )
            
        except Exception as e:
            return EditResult(
                success=False,
                errors=[f"Erreur édition PDF: {str(e)}"]
            )
    
    async def _edit_docx(
        self,
        file_path: Path,
        operations: List[EditOperation]
    ) -> EditResult:
        """Édite un DOCX"""
        if not HAS_DOCX:
            return EditResult(
                success=False,
                errors=["python-docx non installé"]
            )
        
        try:
            doc = DocxDocument(str(file_path))
            
            modified_sections = []
            operations_applied = 0
            warnings = []
            
            for op in operations:
                if op.operation_type == EditOperationType.INSERT:
                    # Format: "paragraph:5" ou "end"
                    if op.target == "end":
                        doc.add_paragraph(str(op.content))
                        modified_sections.append("Ajouté paragraphe à la fin")
                    else:
                        parts = op.target.split(':')
                        if parts[0] == "paragraph":
                            para_num = int(parts[1])
                            if 0 <= para_num < len(doc.paragraphs):
                                # Insertion avant paragraphe
                                p = doc.paragraphs[para_num]
                                new_p = p.insert_paragraph_before(str(op.content))
                                modified_sections.append(f"Inséré avant paragraphe {para_num}")
                            else:
                                doc.add_paragraph(str(op.content))
                                modified_sections.append("Ajouté à la fin")
                    operations_applied += 1
                
                elif op.operation_type == EditOperationType.REPLACE:
                    # Remplace texte dans un paragraphe
                    parts = op.target.split(':')
                    if parts[0] == "paragraph":
                        para_num = int(parts[1])
                        if 0 <= para_num < len(doc.paragraphs):
                            doc.paragraphs[para_num].text = str(op.content)
                            modified_sections.append(f"Remplacé paragraphe {para_num}")
                            operations_applied += 1
                        else:
                            warnings.append(f"Paragraphe {para_num} n'existe pas")
                
                elif op.operation_type == EditOperationType.DELETE:
                    parts = op.target.split(':')
                    if parts[0] == "paragraph":
                        para_num = int(parts[1])
                        if 0 <= para_num < len(doc.paragraphs):
                            # Supprime paragraphe
                            p = doc.paragraphs[para_num]
                            p._element.getparent().remove(p._element)
                            modified_sections.append(f"Supprimé paragraphe {para_num}")
                            operations_applied += 1
                        else:
                            warnings.append(f"Paragraphe {para_num} n'existe pas")
                
                elif op.operation_type == EditOperationType.APPEND:
                    doc.add_paragraph(str(op.content))
                    modified_sections.append("Ajouté contenu à la fin")
                    operations_applied += 1
            
            # Sauvegarde
            doc.save(str(file_path))
            
            return EditResult(
                success=True,
                operations_applied=operations_applied,
                warnings=warnings,
                modified_sections=modified_sections
            )
            
        except Exception as e:
            return EditResult(
                success=False,
                errors=[f"Erreur édition DOCX: {str(e)}"]
            )
    
    async def _edit_excel(
        self,
        file_path: Path,
        operations: List[EditOperation]
    ) -> EditResult:
        """Édite un fichier Excel"""
        if not HAS_EXCEL:
            return EditResult(
                success=False,
                errors=["openpyxl non installé"]
            )
        
        try:
            wb = load_workbook(str(file_path))
            ws = wb.active
            
            modified_sections = []
            operations_applied = 0
            warnings = []
            
            for op in operations:
                if op.operation_type == EditOperationType.UPDATE_CELL:
                    # Format: "cell:A1" ou "cell:B2"
                    cell_ref = op.target.split(':')[1]
                    ws[cell_ref] = op.content
                    modified_sections.append(f"Modifié cellule {cell_ref}")
                    operations_applied += 1
                
                elif op.operation_type == EditOperationType.INSERT:
                    # Insert row/column
                    if "row" in op.target:
                        row_num = int(op.target.split(':')[1])
                        ws.insert_rows(row_num)
                        # Remplit données si fournies
                        if isinstance(op.content, list):
                            for col, value in enumerate(op.content, start=1):
                                ws.cell(row=row_num, column=col, value=value)
                        modified_sections.append(f"Inséré ligne {row_num}")
                        operations_applied += 1
                    elif "column" in op.target:
                        col_num = int(op.target.split(':')[1])
                        ws.insert_cols(col_num)
                        modified_sections.append(f"Inséré colonne {col_num}")
                        operations_applied += 1
                
                elif op.operation_type == EditOperationType.DELETE:
                    if "row" in op.target:
                        row_num = int(op.target.split(':')[1])
                        ws.delete_rows(row_num)
                        modified_sections.append(f"Supprimé ligne {row_num}")
                        operations_applied += 1
                    elif "column" in op.target:
                        col_num = int(op.target.split(':')[1])
                        ws.delete_cols(col_num)
                        modified_sections.append(f"Supprimé colonne {col_num}")
                        operations_applied += 1
                
                elif op.operation_type == EditOperationType.APPEND:
                    # Ajoute ligne à la fin
                    if isinstance(op.content, list):
                        ws.append(op.content)
                        modified_sections.append("Ajouté ligne à la fin")
                        operations_applied += 1
            
            # Sauvegarde
            wb.save(str(file_path))
            
            return EditResult(
                success=True,
                operations_applied=operations_applied,
                warnings=warnings,
                modified_sections=modified_sections
            )
            
        except Exception as e:
            return EditResult(
                success=False,
                errors=[f"Erreur édition Excel: {str(e)}"]
            )
    
    async def _edit_csv(
        self,
        file_path: Path,
        operations: List[EditOperation]
    ) -> EditResult:
        """Édite un fichier CSV"""
        try:
            # Charge CSV
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
            
            modified_sections = []
            operations_applied = 0
            warnings = []
            
            for op in operations:
                if op.operation_type == EditOperationType.INSERT:
                    row_num = int(op.target.split(':')[1]) if ':' in op.target else len(rows)
                    if isinstance(op.content, list):
                        rows.insert(row_num, op.content)
                        modified_sections.append(f"Inséré ligne {row_num}")
                        operations_applied += 1
                
                elif op.operation_type == EditOperationType.DELETE:
                    row_num = int(op.target.split(':')[1])
                    if 0 <= row_num < len(rows):
                        rows.pop(row_num)
                        modified_sections.append(f"Supprimé ligne {row_num}")
                        operations_applied += 1
                    else:
                        warnings.append(f"Ligne {row_num} n'existe pas")
                
                elif op.operation_type == EditOperationType.REPLACE:
                    row_num = int(op.target.split(':')[1])
                    if 0 <= row_num < len(rows):
                        if isinstance(op.content, list):
                            rows[row_num] = op.content
                            modified_sections.append(f"Remplacé ligne {row_num}")
                            operations_applied += 1
                    else:
                        warnings.append(f"Ligne {row_num} n'existe pas")
                
                elif op.operation_type == EditOperationType.UPDATE_CELL:
                    # Format: "cell:2:3" (ligne:colonne)
                    parts = op.target.split(':')
                    row_num = int(parts[1])
                    col_num = int(parts[2])
                    if 0 <= row_num < len(rows) and 0 <= col_num < len(rows[row_num]):
                        rows[row_num][col_num] = str(op.content)
                        modified_sections.append(f"Modifié cellule ({row_num},{col_num})")
                        operations_applied += 1
                    else:
                        warnings.append(f"Cellule ({row_num},{col_num}) invalide")
                
                elif op.operation_type == EditOperationType.APPEND:
                    if isinstance(op.content, list):
                        rows.append(op.content)
                        modified_sections.append("Ajouté ligne à la fin")
                        operations_applied += 1
            
            # Sauvegarde
            with open(file_path, 'w', encoding='utf-8', newline='') as f:
                writer = csv.writer(f)
                writer.writerows(rows)
            
            return EditResult(
                success=True,
                operations_applied=operations_applied,
                warnings=warnings,
                modified_sections=modified_sections
            )
            
        except Exception as e:
            return EditResult(
                success=False,
                errors=[f"Erreur édition CSV: {str(e)}"]
            )
    
    async def _edit_text_file(
        self,
        file_path: Path,
        operations: List[EditOperation]
    ) -> EditResult:
        """Édite un fichier texte (TXT, MD, JSON)"""
        try:
            # Charge contenu
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            
            modified_sections = []
            operations_applied = 0
            warnings = []
            
            for op in operations:
                if op.operation_type == EditOperationType.INSERT:
                    # Format: "line:10" ou "end"
                    if op.target == "end":
                        lines.append(str(op.content))
                        modified_sections.append("Ajouté à la fin")
                    else:
                        line_num = int(op.target.split(':')[1])
                        if 0 <= line_num <= len(lines):
                            lines.insert(line_num, str(op.content))
                            modified_sections.append(f"Inséré ligne {line_num}")
                        else:
                            warnings.append(f"Ligne {line_num} invalide")
                    operations_applied += 1
                
                elif op.operation_type == EditOperationType.DELETE:
                    line_num = int(op.target.split(':')[1])
                    if 0 <= line_num < len(lines):
                        lines.pop(line_num)
                        modified_sections.append(f"Supprimé ligne {line_num}")
                        operations_applied += 1
                    else:
                        warnings.append(f"Ligne {line_num} n'existe pas")
                
                elif op.operation_type == EditOperationType.REPLACE:
                    # Remplacement de texte
                    if ':' in op.target:
                        line_num = int(op.target.split(':')[1])
                        if 0 <= line_num < len(lines):
                            lines[line_num] = str(op.content)
                            modified_sections.append(f"Remplacé ligne {line_num}")
                            operations_applied += 1
                    else:
                        # Remplacement global
                        old_text = op.target
                        new_text = str(op.content)
                        content = content.replace(old_text, new_text)
                        lines = content.split('\n')
                        modified_sections.append(f"Remplacé '{old_text[:30]}...'")
                        operations_applied += 1
                
                elif op.operation_type == EditOperationType.APPEND:
                    lines.append(str(op.content))
                    modified_sections.append("Ajouté à la fin")
                    operations_applied += 1
            
            # Sauvegarde
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))
            
            return EditResult(
                success=True,
                operations_applied=operations_applied,
                warnings=warnings,
                modified_sections=modified_sections
            )
            
        except Exception as e:
            return EditResult(
                success=False,
                errors=[f"Erreur édition fichier texte: {str(e)}"]
            )
    
    async def merge_documents(
        self,
        file_paths: List[str],
        output_path: str,
        format_type: Optional[str] = None
    ) -> EditResult:
        """Fusionne plusieurs documents"""
        
        if not file_paths:
            return EditResult(
                success=False,
                errors=["Aucun fichier à fusionner"]
            )
        
        # Détecte format
        first_file = Path(file_paths[0])
        file_format = format_type or self._detect_format(first_file)
        
        try:
            if file_format == "pdf":
                return await self._merge_pdf(file_paths, output_path)
            elif file_format == "docx":
                return await self._merge_docx(file_paths, output_path)
            elif file_format == "csv":
                return await self._merge_csv(file_paths, output_path)
            elif file_format in ["txt", "md"]:
                return await self._merge_text(file_paths, output_path)
            else:
                return EditResult(
                    success=False,
                    errors=[f"Fusion non supportée pour format {file_format}"]
                )
                
        except Exception as e:
            return EditResult(
                success=False,
                errors=[f"Erreur fusion: {str(e)}"]
            )
    
    async def _merge_pdf(self, file_paths: List[str], output_path: str) -> EditResult:
        """Fusionne plusieurs PDF"""
        if not HAS_PDF:
            return EditResult(success=False, errors=["PyPDF2 non installé"])
        
        writer = PdfWriter()
        
        for file_path in file_paths:
            reader = PdfReader(file_path)
            for page in reader.pages:
                writer.add_page(page)
        
        with open(output_path, 'wb') as f:
            writer.write(f)
        
        return EditResult(
            success=True,
            operations_applied=len(file_paths),
            modified_sections=[f"Fusionné {len(file_paths)} fichiers PDF"]
        )
    
    async def _merge_docx(self, file_paths: List[str], output_path: str) -> EditResult:
        """Fusionne plusieurs DOCX"""
        if not HAS_DOCX:
            return EditResult(success=False, errors=["python-docx non installé"])
        
        merged_doc = DocxDocument()
        
        for file_path in file_paths:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                merged_doc.add_paragraph(para.text)
            merged_doc.add_page_break()
        
        merged_doc.save(output_path)
        
        return EditResult(
            success=True,
            operations_applied=len(file_paths),
            modified_sections=[f"Fusionné {len(file_paths)} fichiers DOCX"]
        )
    
    async def _merge_csv(self, file_paths: List[str], output_path: str) -> EditResult:
        """Fusionne plusieurs CSV"""
        all_rows = []
        headers = None
        
        for file_path in file_paths:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
                if rows:
                    if headers is None:
                        headers = rows[0]
                        all_rows.append(headers)
                    all_rows.extend(rows[1:])  # Skip header des fichiers suivants
        
        with open(output_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)
            writer.writerows(all_rows)
        
        return EditResult(
            success=True,
            operations_applied=len(file_paths),
            modified_sections=[f"Fusionné {len(file_paths)} fichiers CSV"]
        )
    
    async def _merge_text(self, file_paths: List[str], output_path: str) -> EditResult:
        """Fusionne plusieurs fichiers texte"""
        merged_content = []
        
        for file_path in file_paths:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                merged_content.append(content)
                merged_content.append("\n\n---\n\n")  # Séparateur
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("".join(merged_content))
        
        return EditResult(
            success=True,
            operations_applied=len(file_paths),
            modified_sections=[f"Fusionné {len(file_paths)} fichiers texte"]
        )
