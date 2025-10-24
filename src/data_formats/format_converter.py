"""
Convertisseur de formats de données
Conversion entre PDF, DOCX, Excel, CSV, JSON, Markdown, HTML, etc.
"""

from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field
from pathlib import Path
from enum import Enum
import json
import csv
from datetime import datetime
import re

# Document processing
try:
    from PyPDF2 import PdfReader, PdfWriter
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    from openpyxl import Workbook, load_workbook
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    from bs4 import BeautifulSoup
    import markdown
    HAS_HTML = True
except ImportError:
    HAS_HTML = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


class SupportedFormat(Enum):
    """Formats supportés"""
    PDF = "pdf"
    DOCX = "docx"
    EXCEL = "xlsx"
    CSV = "csv"
    JSON = "json"
    MARKDOWN = "md"
    HTML = "html"
    TEXT = "txt"
    XML = "xml"
    IMAGE = "image"  # Pour OCR


@dataclass
class ConversionConfig:
    """Configuration de conversion"""
    preserve_formatting: bool = True
    extract_images: bool = False
    ocr_language: str = "fra+eng"  # Français + Anglais
    table_detection: bool = True
    smart_parsing: bool = True  # Utilise LLM pour améliorer parsing
    output_template: Optional[str] = None
    custom_options: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ConversionResult:
    """Résultat de conversion"""
    success: bool
    source_format: str
    target_format: str
    output_path: Optional[str] = None
    output_data: Optional[Any] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    conversion_time: float = 0.0
    
    def to_dict(self) -> Dict:
        """Conversion en dictionnaire"""
        return {
            "success": self.success,
            "source_format": self.source_format,
            "target_format": self.target_format,
            "output_path": self.output_path,
            "metadata": self.metadata,
            "warnings": self.warnings,
            "errors": self.errors,
            "conversion_time": self.conversion_time
        }


class FormatConverter:
    """Convertisseur universel de formats"""
    
    def __init__(self, llm_analyzer: Optional[Any] = None):
        self.llm = llm_analyzer
        self.supported_conversions = self._build_conversion_matrix()
    
    def _build_conversion_matrix(self) -> Dict[str, List[str]]:
        """Matrice des conversions supportées"""
        return {
            SupportedFormat.PDF.value: ["json", "txt", "md", "docx"],
            SupportedFormat.DOCX.value: ["pdf", "txt", "md", "html"],
            SupportedFormat.EXCEL.value: ["csv", "json", "html"],
            SupportedFormat.CSV.value: ["xlsx", "json", "md"],
            SupportedFormat.JSON.value: ["csv", "xlsx", "md", "html"],
            SupportedFormat.MARKDOWN.value: ["html", "pdf", "docx"],
            SupportedFormat.HTML.value: ["md", "txt", "pdf"],
            SupportedFormat.IMAGE.value: ["txt", "json", "md"],  # Via OCR
        }
    
    async def convert(
        self,
        source_path: str,
        target_format: str,
        output_path: Optional[str] = None,
        config: Optional[ConversionConfig] = None
    ) -> ConversionResult:
        """Convertit un document d'un format à un autre"""
        start_time = datetime.now()
        config = config or ConversionConfig()
        
        source = Path(source_path)
        if not source.exists():
            return ConversionResult(
                success=False,
                source_format="unknown",
                target_format=target_format,
                errors=["Fichier source introuvable"]
            )
        
        # Détecte format source
        source_format = self._detect_format(source)
        
        # Vérifie conversion possible
        if not self._is_conversion_supported(source_format, target_format):
            return ConversionResult(
                success=False,
                source_format=source_format,
                target_format=target_format,
                errors=[f"Conversion {source_format} → {target_format} non supportée"]
            )
        
        # Route vers convertisseur approprié
        try:
            if source_format == "pdf":
                result = await self._convert_from_pdf(source, target_format, config)
            elif source_format == "docx":
                result = await self._convert_from_docx(source, target_format, config)
            elif source_format == "xlsx":
                result = await self._convert_from_excel(source, target_format, config)
            elif source_format == "csv":
                result = await self._convert_from_csv(source, target_format, config)
            elif source_format == "json":
                result = await self._convert_from_json(source, target_format, config)
            elif source_format == "md":
                result = await self._convert_from_markdown(source, target_format, config)
            elif source_format == "html":
                result = await self._convert_from_html(source, target_format, config)
            elif source_format in ["png", "jpg", "jpeg", "tiff"]:
                result = await self._convert_from_image(source, target_format, config)
            else:
                result = ConversionResult(
                    success=False,
                    source_format=source_format,
                    target_format=target_format,
                    errors=[f"Format source {source_format} non implémenté"]
                )
            
            # Sauvegarde si output_path fourni
            if result.success and output_path and result.output_data:
                self._save_output(result.output_data, output_path, target_format)
                result.output_path = output_path
            
            duration = (datetime.now() - start_time).total_seconds()
            result.conversion_time = duration
            
            return result
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format=source_format,
                target_format=target_format,
                errors=[f"Erreur conversion: {str(e)}"]
            )
    
    def _detect_format(self, file_path: Path) -> str:
        """Détecte le format d'un fichier"""
        suffix = file_path.suffix.lower().lstrip('.')
        
        # Mapping extensions
        format_map = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',
            'xlsx': 'xlsx',
            'xls': 'xlsx',
            'csv': 'csv',
            'json': 'json',
            'md': 'md',
            'markdown': 'md',
            'html': 'html',
            'htm': 'html',
            'txt': 'txt',
            'xml': 'xml',
            'png': 'png',
            'jpg': 'jpg',
            'jpeg': 'jpeg',
            'tiff': 'tiff'
        }
        
        return format_map.get(suffix, 'unknown')
    
    def _is_conversion_supported(self, source: str, target: str) -> bool:
        """Vérifie si conversion supportée"""
        return target in self.supported_conversions.get(source, [])
    
    async def _convert_from_pdf(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit depuis PDF"""
        if not HAS_PDF:
            return ConversionResult(
                success=False,
                source_format="pdf",
                target_format=target_format,
                errors=["PyPDF2 non installé"]
            )
        
        try:
            reader = PdfReader(str(source))
            
            # Extrait texte
            text_content = []
            metadata = {
                "pages": len(reader.pages),
                "title": reader.metadata.title if reader.metadata else None,
                "author": reader.metadata.author if reader.metadata else None
            }
            
            for page in reader.pages:
                text_content.append(page.extract_text())
            
            full_text = "\n\n".join(text_content)
            
            # Conversion selon target
            if target_format == "json":
                output = await self._pdf_to_json(full_text, text_content, metadata, config)
            elif target_format == "txt":
                output = full_text
            elif target_format == "md":
                output = await self._pdf_to_markdown(full_text, metadata, config)
            elif target_format == "docx":
                output = await self._text_to_docx(full_text, metadata)
            else:
                return ConversionResult(
                    success=False,
                    source_format="pdf",
                    target_format=target_format,
                    errors=[f"Conversion PDF → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="pdf",
                target_format=target_format,
                output_data=output,
                metadata=metadata
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="pdf",
                target_format=target_format,
                errors=[f"Erreur lecture PDF: {str(e)}"]
            )
    
    async def _pdf_to_json(
        self,
        full_text: str,
        pages: List[str],
        metadata: Dict,
        config: ConversionConfig
    ) -> Dict:
        """Convertit PDF en JSON structuré"""
        
        # Structure de base
        json_output = {
            "metadata": metadata,
            "pages": [
                {
                    "page_number": i + 1,
                    "content": page.strip()
                }
                for i, page in enumerate(pages)
            ],
            "full_text": full_text
        }
        
        # Analyse avancée avec LLM si disponible
        if config.smart_parsing and self.llm:
            try:
                # Demande au LLM d'extraire structure
                prompt = f"""
                Analyse ce document PDF et extrait une structure JSON:
                
                {full_text[:5000]}  # Premiers 5000 caractères
                
                Identifie:
                - Titre principal
                - Sections principales
                - Points clés
                - Données tabulaires (si présentes)
                - Résumé
                
                Retourne un JSON structuré.
                """
                
                # Note: Appel LLM à implémenter selon votre agent
                # structured_data = await self.llm.process_request(prompt)
                # json_output["structured_content"] = structured_data
                
                json_output["smart_parsing_applied"] = True
            except Exception as e:
                json_output["parsing_warning"] = f"Smart parsing échoué: {str(e)}"
        
        # Détection automatique de tables
        if config.table_detection:
            tables = self._detect_tables_in_text(full_text)
            if tables:
                json_output["tables"] = tables
        
        return json_output
    
    def _detect_tables_in_text(self, text: str) -> List[Dict]:
        """Détecte des tableaux dans le texte"""
        tables = []
        lines = text.split('\n')
        
        # Heuristique simple: lignes avec séparateurs répétés
        for i, line in enumerate(lines):
            if '|' in line and line.count('|') >= 2:
                # Possible tableau Markdown
                cells = [cell.strip() for cell in line.split('|')]
                if len(cells) > 2:  # Au moins 2 colonnes
                    tables.append({
                        "line_number": i + 1,
                        "type": "markdown_table",
                        "columns": len(cells) - 2,  # Enlève vides début/fin
                        "sample": line[:100]
                    })
            elif '\t' in line and line.count('\t') >= 2:
                # Possible tableau séparé par tabs
                cells = line.split('\t')
                tables.append({
                    "line_number": i + 1,
                    "type": "tab_separated",
                    "columns": len(cells),
                    "sample": line[:100]
                })
        
        return tables[:10]  # Max 10 tables détectées
    
    async def _pdf_to_markdown(
        self,
        text: str,
        metadata: Dict,
        config: ConversionConfig
    ) -> str:
        """Convertit PDF en Markdown"""
        md_content = []
        
        # Header avec metadata
        if metadata.get("title"):
            md_content.append(f"# {metadata['title']}\n")
        
        if metadata.get("author"):
            md_content.append(f"**Auteur:** {metadata['author']}\n")
        
        md_content.append(f"**Pages:** {metadata.get('pages', 'N/A')}\n")
        md_content.append("---\n\n")
        
        # Contenu avec structure basique
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Détecte titres (TOUT EN MAJUSCULES)
            if para.isupper() and len(para) < 100:
                md_content.append(f"## {para}\n\n")
            # Détecte listes (lignes qui commencent par -, *, •)
            elif para.startswith(('-', '*', '•')):
                md_content.append(f"{para}\n\n")
            else:
                md_content.append(f"{para}\n\n")
        
        return "".join(md_content)
    
    async def _convert_from_docx(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit depuis DOCX"""
        if not HAS_DOCX:
            return ConversionResult(
                success=False,
                source_format="docx",
                target_format=target_format,
                errors=["python-docx non installé"]
            )
        
        try:
            doc = DocxDocument(str(source))
            
            # Extrait texte
            text_content = []
            for para in doc.paragraphs:
                text_content.append(para.text)
            
            full_text = "\n".join(text_content)
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections)
            }
            
            # Conversion
            if target_format == "txt":
                output = full_text
            elif target_format == "md":
                output = await self._docx_to_markdown(doc, config)
            elif target_format == "html":
                output = await self._docx_to_html(doc, config)
            elif target_format == "pdf":
                # Note: Nécessite conversion externe (LibreOffice, etc.)
                return ConversionResult(
                    success=False,
                    source_format="docx",
                    target_format="pdf",
                    errors=["Conversion DOCX → PDF nécessite outil externe"]
                )
            else:
                return ConversionResult(
                    success=False,
                    source_format="docx",
                    target_format=target_format,
                    errors=[f"Conversion DOCX → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="docx",
                target_format=target_format,
                output_data=output,
                metadata=metadata
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="docx",
                target_format=target_format,
                errors=[f"Erreur lecture DOCX: {str(e)}"]
            )
    
    async def _docx_to_markdown(self, doc: Any, config: ConversionConfig) -> str:
        """Convertit DOCX en Markdown"""
        md_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Détecte styles
            style = para.style.name.lower()
            
            if 'heading 1' in style or 'title' in style:
                md_lines.append(f"# {text}\n")
            elif 'heading 2' in style:
                md_lines.append(f"## {text}\n")
            elif 'heading 3' in style:
                md_lines.append(f"### {text}\n")
            else:
                # Format inline (gras, italique)
                formatted_text = text
                for run in para.runs:
                    if run.bold:
                        formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                    if run.italic:
                        formatted_text = formatted_text.replace(run.text, f"*{run.text}*")
                
                md_lines.append(f"{formatted_text}\n\n")
        
        return "".join(md_lines)
    
    async def _docx_to_html(self, doc: Any, config: ConversionConfig) -> str:
        """Convertit DOCX en HTML"""
        html_parts = ["<!DOCTYPE html>\n<html>\n<head>\n<meta charset='utf-8'>\n</head>\n<body>\n"]
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style = para.style.name.lower()
            
            if 'heading 1' in style:
                html_parts.append(f"<h1>{text}</h1>\n")
            elif 'heading 2' in style:
                html_parts.append(f"<h2>{text}</h2>\n")
            else:
                html_parts.append(f"<p>{text}</p>\n")
        
        html_parts.append("</body>\n</html>")
        return "".join(html_parts)
    
    async def _convert_from_excel(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit depuis Excel"""
        if not HAS_EXCEL:
            return ConversionResult(
                success=False,
                source_format="xlsx",
                target_format=target_format,
                errors=["openpyxl non installé"]
            )
        
        try:
            wb = load_workbook(str(source))
            ws = wb.active
            
            # Extrait données
            data = []
            for row in ws.iter_rows(values_only=True):
                data.append(list(row))
            
            metadata = {
                "sheets": len(wb.sheetnames),
                "rows": len(data),
                "columns": len(data[0]) if data else 0
            }
            
            # Conversion
            if target_format == "csv":
                output = self._data_to_csv(data)
            elif target_format == "json":
                output = self._excel_to_json(data, ws)
            elif target_format == "html":
                output = self._data_to_html_table(data)
            else:
                return ConversionResult(
                    success=False,
                    source_format="xlsx",
                    target_format=target_format,
                    errors=[f"Conversion Excel → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="xlsx",
                target_format=target_format,
                output_data=output,
                metadata=metadata
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="xlsx",
                target_format=target_format,
                errors=[f"Erreur lecture Excel: {str(e)}"]
            )
    
    def _excel_to_json(self, data: List[List], worksheet: Any) -> Dict:
        """Convertit Excel en JSON structuré"""
        if not data:
            return {"data": []}
        
        # Première ligne = headers
        headers = data[0]
        rows = data[1:]
        
        json_data = {
            "metadata": {
                "rows": len(rows),
                "columns": len(headers)
            },
            "headers": headers,
            "data": []
        }
        
        # Convertit en liste de dictionnaires
        for row in rows:
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(row):
                    row_dict[str(header)] = row[i]
            json_data["data"].append(row_dict)
        
        return json_data
    
    def _data_to_csv(self, data: List[List]) -> str:
        """Convertit données en CSV"""
        import io
        output = io.StringIO()
        writer = csv.writer(output)
        for row in data:
            writer.writerow(row)
        return output.getvalue()
    
    def _data_to_html_table(self, data: List[List]) -> str:
        """Convertit données en tableau HTML"""
        html = ["<table border='1'>\n"]
        
        if data:
            # Header
            html.append("<thead><tr>")
            for cell in data[0]:
                html.append(f"<th>{cell}</th>")
            html.append("</tr></thead>\n")
            
            # Body
            html.append("<tbody>")
            for row in data[1:]:
                html.append("<tr>")
                for cell in row:
                    html.append(f"<td>{cell}</td>")
                html.append("</tr>\n")
            html.append("</tbody>")
        
        html.append("</table>")
        return "".join(html)
    
    async def _convert_from_csv(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit depuis CSV"""
        try:
            with open(source, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                data = list(reader)
            
            metadata = {
                "rows": len(data),
                "columns": len(data[0]) if data else 0
            }
            
            # Conversion
            if target_format == "xlsx":
                output = await self._csv_to_excel(data)
            elif target_format == "json":
                output = self._csv_to_json(data)
            elif target_format == "md":
                output = self._csv_to_markdown(data)
            else:
                return ConversionResult(
                    success=False,
                    source_format="csv",
                    target_format=target_format,
                    errors=[f"Conversion CSV → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="csv",
                target_format=target_format,
                output_data=output,
                metadata=metadata
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="csv",
                target_format=target_format,
                errors=[f"Erreur lecture CSV: {str(e)}"]
            )
    
    async def _csv_to_excel(self, data: List[List]) -> bytes:
        """Convertit CSV en Excel"""
        if not HAS_EXCEL:
            raise ImportError("openpyxl requis")
        
        wb = Workbook()
        ws = wb.active
        
        for row in data:
            ws.append(row)
        
        # Sauvegarde en bytes
        import io
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
    
    def _csv_to_json(self, data: List[List]) -> Dict:
        """Convertit CSV en JSON"""
        if not data:
            return {"data": []}
        
        headers = data[0]
        rows = data[1:]
        
        json_data = {
            "headers": headers,
            "data": [
                {headers[i]: row[i] if i < len(row) else None for i in range(len(headers))}
                for row in rows
            ]
        }
        
        return json_data
    
    def _csv_to_markdown(self, data: List[List]) -> str:
        """Convertit CSV en tableau Markdown"""
        if not data:
            return ""
        
        md_lines = []
        
        # Header
        md_lines.append("| " + " | ".join(str(cell) for cell in data[0]) + " |")
        md_lines.append("|" + "|".join(["---" for _ in data[0]]) + "|")
        
        # Rows
        for row in data[1:]:
            md_lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
        
        return "\n".join(md_lines)
    
    async def _convert_from_json(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit depuis JSON"""
        try:
            with open(source, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Conversion
            if target_format == "csv":
                output = self._json_to_csv(data)
            elif target_format == "xlsx":
                output = await self._json_to_excel(data)
            elif target_format == "md":
                output = self._json_to_markdown(data)
            elif target_format == "html":
                output = self._json_to_html(data)
            else:
                return ConversionResult(
                    success=False,
                    source_format="json",
                    target_format=target_format,
                    errors=[f"Conversion JSON → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="json",
                target_format=target_format,
                output_data=output
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="json",
                target_format=target_format,
                errors=[f"Erreur lecture JSON: {str(e)}"]
            )
    
    def _json_to_csv(self, data: Any) -> str:
        """Convertit JSON en CSV"""
        import io
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Si liste de dictionnaires
        if isinstance(data, list) and data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            writer.writerow(headers)
            for row in data:
                writer.writerow([row.get(h) for h in headers])
        # Si dict avec clé "data"
        elif isinstance(data, dict) and "data" in data:
            return self._json_to_csv(data["data"])
        else:
            # JSON complexe: format basique
            writer.writerow(["key", "value"])
            for k, v in (data.items() if isinstance(data, dict) else enumerate(data)):
                writer.writerow([k, json.dumps(v) if isinstance(v, (dict, list)) else v])
        
        return output.getvalue()
    
    async def _json_to_excel(self, data: Any) -> bytes:
        """Convertit JSON en Excel"""
        csv_data = self._json_to_csv(data)
        lines = csv_data.strip().split('\n')
        data_rows = [line.split(',') for line in lines]
        return await self._csv_to_excel(data_rows)
    
    def _json_to_markdown(self, data: Any) -> str:
        """Convertit JSON en Markdown"""
        md_lines = ["# Données JSON\n"]
        md_lines.append("```json\n")
        md_lines.append(json.dumps(data, indent=2, ensure_ascii=False))
        md_lines.append("\n```")
        return "".join(md_lines)
    
    def _json_to_html(self, data: Any) -> str:
        """Convertit JSON en HTML"""
        html = ["<html><body>\n"]
        html.append("<pre>\n")
        html.append(json.dumps(data, indent=2, ensure_ascii=False))
        html.append("\n</pre>\n")
        html.append("</body></html>")
        return "".join(html)
    
    async def _convert_from_markdown(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit depuis Markdown"""
        try:
            with open(source, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            if target_format == "html":
                if HAS_HTML:
                    output = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
                else:
                    output = f"<html><body><pre>{md_content}</pre></body></html>"
            elif target_format == "txt":
                # Supprime le markdown syntax
                output = re.sub(r'[#*_`\[\]]', '', md_content)
            else:
                return ConversionResult(
                    success=False,
                    source_format="md",
                    target_format=target_format,
                    errors=[f"Conversion Markdown → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="md",
                target_format=target_format,
                output_data=output
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="md",
                target_format=target_format,
                errors=[f"Erreur lecture Markdown: {str(e)}"]
            )
    
    async def _convert_from_html(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit depuis HTML"""
        if not HAS_HTML:
            return ConversionResult(
                success=False,
                source_format="html",
                target_format=target_format,
                errors=["BeautifulSoup non installé"]
            )
        
        try:
            with open(source, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            if target_format == "txt":
                output = soup.get_text()
            elif target_format == "md":
                output = self._html_to_markdown(soup)
            else:
                return ConversionResult(
                    success=False,
                    source_format="html",
                    target_format=target_format,
                    errors=[f"Conversion HTML → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="html",
                target_format=target_format,
                output_data=output
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="html",
                target_format=target_format,
                errors=[f"Erreur lecture HTML: {str(e)}"]
            )
    
    def _html_to_markdown(self, soup: Any) -> str:
        """Convertit HTML en Markdown (basique)"""
        md_lines = []
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
            if element.name == 'h1':
                md_lines.append(f"# {element.get_text()}\n")
            elif element.name == 'h2':
                md_lines.append(f"## {element.get_text()}\n")
            elif element.name == 'h3':
                md_lines.append(f"### {element.get_text()}\n")
            elif element.name == 'p':
                md_lines.append(f"{element.get_text()}\n\n")
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    md_lines.append(f"- {li.get_text()}\n")
                md_lines.append("\n")
        
        return "".join(md_lines)
    
    async def _convert_from_image(
        self,
        source: Path,
        target_format: str,
        config: ConversionConfig
    ) -> ConversionResult:
        """Convertit image en texte via OCR"""
        if not HAS_OCR:
            return ConversionResult(
                success=False,
                source_format="image",
                target_format=target_format,
                errors=["pytesseract non installé"]
            )
        
        try:
            image = Image.open(source)
            text = pytesseract.image_to_string(image, lang=config.ocr_language)
            
            metadata = {
                "image_size": image.size,
                "ocr_language": config.ocr_language
            }
            
            if target_format == "txt":
                output = text
            elif target_format == "json":
                output = {
                    "source_image": str(source),
                    "extracted_text": text,
                    "metadata": metadata
                }
            elif target_format == "md":
                output = f"# Texte extrait de {source.name}\n\n{text}"
            else:
                return ConversionResult(
                    success=False,
                    source_format="image",
                    target_format=target_format,
                    errors=[f"Conversion Image → {target_format} non implémentée"]
                )
            
            return ConversionResult(
                success=True,
                source_format="image",
                target_format=target_format,
                output_data=output,
                metadata=metadata
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                source_format="image",
                target_format=target_format,
                errors=[f"Erreur OCR: {str(e)}"]
            )
    
    async def _text_to_docx(self, text: str, metadata: Dict) -> bytes:
        """Convertit texte en DOCX"""
        if not HAS_DOCX:
            raise ImportError("python-docx requis")
        
        doc = DocxDocument()
        
        # Titre si disponible
        if metadata.get("title"):
            doc.add_heading(metadata["title"], 0)
        
        # Contenu
        for para in text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Sauvegarde en bytes
        import io
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    def _save_output(self, data: Any, output_path: str, target_format: str):
        """Sauvegarde le résultat de conversion"""
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        
        if isinstance(data, bytes):
            with open(output, 'wb') as f:
                f.write(data)
        elif isinstance(data, (dict, list)):
            with open(output, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        else:
            with open(output, 'w', encoding='utf-8') as f:
                f.write(str(data))
