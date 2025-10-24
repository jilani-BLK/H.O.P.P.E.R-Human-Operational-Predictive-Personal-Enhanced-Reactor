"""
Générateur de documents HOPPER
Crée des documents professionnels à partir de données
"""

import json
import yaml
from pathlib import Path
from typing import Dict, Any, Optional, Union, List
from dataclasses import dataclass
from datetime import datetime

# Imports optionnels
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Backend sans GUI
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    from jinja2 import Environment, FileSystemLoader, Template
    HAS_JINJA2 = True
except ImportError:
    HAS_JINJA2 = False


@dataclass
class DocumentTemplate:
    """Template de document"""
    name: str
    format: str
    content: str
    variables: List[str]
    description: Optional[str] = None


@dataclass
class GenerationResult:
    """Résultat de la génération de document"""
    success: bool
    output_path: Optional[Path] = None
    format: Optional[str] = None
    error: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


@dataclass
class GenerationConfig:
    """Configuration pour la génération de documents"""
    page_size: str = "A4"  # A4, letter, etc.
    orientation: str = "portrait"  # portrait, landscape
    margins: Optional[Dict[str, float]] = None  # top, bottom, left, right en cm
    font_family: str = "Helvetica"
    font_size: int = 11
    include_header: bool = True
    include_footer: bool = True
    include_toc: bool = False  # Table des matières
    include_page_numbers: bool = True
    style: str = "professional"  # professional, modern, minimal, colorful
    
    def __post_init__(self):
        if self.margins is None:
            self.margins = {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5}


class DocumentGenerator:
    """Générateur de documents à partir de données"""
    
    def __init__(self):
        self.templates_dir = Path("templates")
        self.templates_dir.mkdir(exist_ok=True)
        
        # Styles prédéfinis
        self.styles = {
            "professional": {
                "primary_color": (0, 51, 102),  # Bleu foncé
                "secondary_color": (100, 149, 237),  # Bleu clair
                "accent_color": (255, 165, 0),  # Orange
                "text_color": (51, 51, 51),  # Gris foncé
                "bg_color": (255, 255, 255)  # Blanc
            },
            "modern": {
                "primary_color": (63, 81, 181),  # Indigo
                "secondary_color": (156, 39, 176),  # Violet
                "accent_color": (255, 193, 7),  # Ambre
                "text_color": (33, 33, 33),
                "bg_color": (250, 250, 250)
            },
            "minimal": {
                "primary_color": (33, 33, 33),
                "secondary_color": (97, 97, 97),
                "accent_color": (189, 189, 189),
                "text_color": (66, 66, 66),
                "bg_color": (255, 255, 255)
            },
            "colorful": {
                "primary_color": (244, 67, 54),  # Rouge
                "secondary_color": (33, 150, 243),  # Bleu
                "accent_color": (76, 175, 80),  # Vert
                "text_color": (33, 33, 33),
                "bg_color": (255, 255, 255)
            }
        }
    
    async def generate_pdf_from_data(
        self,
        data: Dict[str, Any],
        output: Union[str, Path],
        template: Optional[str] = None,
        config: Optional[GenerationConfig] = None
    ) -> Path:
        """
        Génère un PDF à partir de données structurées
        
        Args:
            data: Données au format dict
            output: Chemin du fichier de sortie
            template: Template Jinja2 optionnel
            config: Configuration de génération
        """
        if not HAS_REPORTLAB:
            raise ImportError("reportlab requis pour PDF. pip install reportlab")
        
        config = config or GenerationConfig()
        output = Path(output)
        
        # Créer le document
        page_size = A4 if config.page_size == "A4" else letter
        doc = SimpleDocTemplate(
            str(output),
            pagesize=page_size,
            rightMargin=config.margins["right"] * cm,
            leftMargin=config.margins["left"] * cm,
            topMargin=config.margins["top"] * cm,
            bottomMargin=config.margins["bottom"] * cm
        )
        
        # Styles
        styles = getSampleStyleSheet()
        style_colors = self.styles.get(config.style, self.styles["professional"])
        
        # Créer style personnalisé
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.Color(
                style_colors["primary_color"][0] / 255,
                style_colors["primary_color"][1] / 255,
                style_colors["primary_color"][2] / 255
            ),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.Color(
                style_colors["secondary_color"][0] / 255,
                style_colors["secondary_color"][1] / 255,
                style_colors["secondary_color"][2] / 255
            ),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=config.font_size,
            textColor=colors.Color(
                style_colors["text_color"][0] / 255,
                style_colors["text_color"][1] / 255,
                style_colors["text_color"][2] / 255
            ),
            fontName=config.font_family
        )
        
        # Construire le contenu
        story = []
        
        # Titre principal
        if "title" in data:
            story.append(Paragraph(data["title"], title_style))
            story.append(Spacer(1, 0.5 * inch))
        
        # Métadonnées
        if "metadata" in data:
            meta = data["metadata"]
            for key, value in meta.items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", normal_style))
            story.append(Spacer(1, 0.3 * inch))
        
        # Sections
        if "sections" in data:
            for section in data["sections"]:
                # Titre de section
                if "title" in section:
                    story.append(Paragraph(section["title"], heading_style))
                    story.append(Spacer(1, 0.2 * inch))
                
                # Contenu
                if "content" in section:
                    if isinstance(section["content"], str):
                        story.append(Paragraph(section["content"], normal_style))
                    elif isinstance(section["content"], list):
                        for para in section["content"]:
                            story.append(Paragraph(str(para), normal_style))
                            story.append(Spacer(1, 0.1 * inch))
                    story.append(Spacer(1, 0.2 * inch))
                
                # Tableau
                if "table" in section:
                    table_data = section["table"]
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(
                            style_colors["primary_color"][0] / 255,
                            style_colors["primary_color"][1] / 255,
                            style_colors["primary_color"][2] / 255
                        )),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.3 * inch))
                
                # Image
                if "image" in section and Path(section["image"]).exists():
                    img = RLImage(section["image"], width=4 * inch, height=3 * inch)
                    story.append(img)
                    story.append(Spacer(1, 0.2 * inch))
                
                # Saut de page optionnel
                if section.get("page_break", False):
                    story.append(PageBreak())
        
        # Liste
        if "list" in data:
            for item in data["list"]:
                story.append(Paragraph(f"• {item}", normal_style))
            story.append(Spacer(1, 0.2 * inch))
        
        # Pied de page avec date
        if config.include_footer:
            story.append(Spacer(1, 0.5 * inch))
            footer_text = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            footer_style = ParagraphStyle(
                'Footer',
                parent=normal_style,
                fontSize=8,
                textColor=colors.grey,
                alignment=TA_CENTER
            )
            story.append(Paragraph(footer_text, footer_style))
        
        # Générer le PDF
        doc.build(story)
        
        return output
    
    async def generate_docx_from_data(
        self,
        data: Dict[str, Any],
        output: Union[str, Path],
        template: Optional[str] = None,
        config: Optional[GenerationConfig] = None
    ) -> Path:
        """Génère un document Word à partir de données"""
        if not HAS_DOCX:
            raise ImportError("python-docx requis. pip install python-docx")
        
        config = config or GenerationConfig()
        output = Path(output)
        
        # Créer le document
        doc = Document()
        style_colors = self.styles.get(config.style, self.styles["professional"])
        
        # Titre principal
        if "title" in data:
            title = doc.add_heading(data["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(*style_colors["primary_color"])
        
        # Métadonnées
        if "metadata" in data:
            for key, value in data["metadata"].items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
        
        # Sections
        if "sections" in data:
            for section in data["sections"]:
                # Titre de section
                if "title" in section:
                    heading = doc.add_heading(section["title"], level=1)
                    heading_run = heading.runs[0]
                    heading_run.font.color.rgb = RGBColor(*style_colors["secondary_color"])
                
                # Contenu
                if "content" in section:
                    if isinstance(section["content"], str):
                        doc.add_paragraph(section["content"])
                    elif isinstance(section["content"], list):
                        for para in section["content"]:
                            doc.add_paragraph(str(para))
                
                # Tableau
                if "table" in section:
                    table_data = section["table"]
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row in enumerate(table_data):
                        for j, cell_value in enumerate(row):
                            cell = table.rows[i].cells[j]
                            cell.text = str(cell_value)
                            
                            # En-tête en gras
                            if i == 0:
                                cell.paragraphs[0].runs[0].font.bold = True
                
                # Image
                if "image" in section and Path(section["image"]).exists():
                    doc.add_picture(section["image"], width=Inches(5))
                
                # Saut de page
                if section.get("page_break", False):
                    doc.add_page_break()
        
        # Liste
        if "list" in data:
            for item in data["list"]:
                doc.add_paragraph(item, style='List Bullet')
        
        # Pied de page
        if config.include_footer:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sauvegarder
        doc.save(str(output))
        
        return output
    
    async def generate_excel_from_data(
        self,
        data: Dict[str, Any],
        output: Union[str, Path],
        include_charts: bool = True,
        config: Optional[GenerationConfig] = None
    ) -> Path:
        """Génère un fichier Excel avec données et graphiques"""
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl requis. pip install openpyxl pandas")
        
        config = config or GenerationConfig()
        output = Path(output)
        style_colors = self.styles.get(config.style, self.styles["professional"])
        
        wb = Workbook()
        if wb.active:
            wb.remove(wb.active)  # Retirer la feuille par défaut
        
        # Créer les feuilles à partir des données
        if "sheets" in data:
            for sheet_data in data["sheets"]:
                sheet_name = sheet_data.get("name", "Feuille")
                ws = wb.create_sheet(title=sheet_name)
                
                # En-têtes
                if "headers" in sheet_data:
                    headers = sheet_data["headers"]
                    for col_idx, header in enumerate(headers, 1):
                        cell = ws.cell(row=1, column=col_idx, value=header)
                        cell.font = Font(bold=True, color="FFFFFF")
                        cell.fill = PatternFill(
                            start_color=f"{style_colors['primary_color'][0]:02x}"
                                      f"{style_colors['primary_color'][1]:02x}"
                                      f"{style_colors['primary_color'][2]:02x}",
                            end_color=f"{style_colors['primary_color'][0]:02x}"
                                    f"{style_colors['primary_color'][1]:02x}"
                                    f"{style_colors['primary_color'][2]:02x}",
                            fill_type="solid"
                        )
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Données
                if "data" in sheet_data:
                    for row_idx, row in enumerate(sheet_data["data"], 2):
                        for col_idx, value in enumerate(row, 1):
                            ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Ajuster largeur des colonnes
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
                
                # Graphiques
                if include_charts and "chart" in sheet_data:
                    chart_config = sheet_data["chart"]
                    chart_type = chart_config.get("type", "bar")
                    
                    if chart_type == "bar":
                        chart = BarChart()
                    elif chart_type == "line":
                        chart = LineChart()
                    elif chart_type == "pie":
                        chart = PieChart()
                    else:
                        chart = BarChart()
                    
                    chart.title = chart_config.get("title", "Graphique")
                    chart.style = 10
                    
                    # Données du graphique
                    data_ref = Reference(
                        ws,
                        min_col=2,
                        min_row=1,
                        max_row=len(sheet_data["data"]) + 1,
                        max_col=len(sheet_data["headers"])
                    )
                    cats = Reference(
                        ws,
                        min_col=1,
                        min_row=2,
                        max_row=len(sheet_data["data"]) + 1
                    )
                    
                    chart.add_data(data_ref, titles_from_data=True)
                    chart.set_categories(cats)
                    
                    # Position du graphique
                    ws.add_chart(chart, "A" + str(len(sheet_data["data"]) + 4))
        
        # Sauvegarder
        wb.save(str(output))
        
        return output
    
    async def generate_html_from_data(
        self,
        data: Dict[str, Any],
        output: Union[str, Path],
        template: Optional[str] = None,
        config: Optional[GenerationConfig] = None
    ) -> Path:
        """Génère un fichier HTML à partir de données"""
        config = config or GenerationConfig()
        output = Path(output)
        style_colors = self.styles.get(config.style, self.styles["professional"])
        
        # Template Jinja2 si fourni
        if template and HAS_JINJA2:
            if Path(template).exists():
                env = Environment(loader=FileSystemLoader(Path(template).parent))
                tmpl = env.get_template(Path(template).name)
            else:
                tmpl = Template(template)
            
            html_content = tmpl.render(**data)
        else:
            # Génération HTML basique
            html_parts = [
                "<!DOCTYPE html>",
                "<html lang='fr'>",
                "<head>",
                "  <meta charset='UTF-8'>",
                "  <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
                f"  <title>{data.get('title', 'Document')}</title>",
                "  <style>",
                "    body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }",
                f"    h1 {{ color: rgb{style_colors['primary_color']}; text-align: center; }}",
                f"    h2 {{ color: rgb{style_colors['secondary_color']}; border-bottom: 2px solid rgb{style_colors['accent_color']}; padding-bottom: 5px; }}",
                "    table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
                f"    th {{ background-color: rgb{style_colors['primary_color']}; color: white; padding: 12px; text-align: left; }}",
                "    td { border: 1px solid #ddd; padding: 8px; }",
                "    tr:nth-child(even) { background-color: #f2f2f2; }",
                "    img { max-width: 100%; height: auto; }",
                "    .footer { text-align: center; font-size: 0.9em; color: #666; margin-top: 40px; }",
                "  </style>",
                "</head>",
                "<body>"
            ]
            
            # Titre
            if "title" in data:
                html_parts.append(f"  <h1>{data['title']}</h1>")
            
            # Métadonnées
            if "metadata" in data:
                for key, value in data["metadata"].items():
                    html_parts.append(f"  <p><strong>{key}:</strong> {value}</p>")
            
            # Sections
            if "sections" in data:
                for section in data["sections"]:
                    if "title" in section:
                        html_parts.append(f"  <h2>{section['title']}</h2>")
                    
                    if "content" in section:
                        if isinstance(section["content"], str):
                            html_parts.append(f"  <p>{section['content']}</p>")
                        elif isinstance(section["content"], list):
                            for para in section["content"]:
                                html_parts.append(f"  <p>{para}</p>")
                    
                    if "table" in section:
                        html_parts.append("  <table>")
                        for i, row in enumerate(section["table"]):
                            html_parts.append("    <tr>")
                            tag = "th" if i == 0 else "td"
                            for cell in row:
                                html_parts.append(f"      <{tag}>{cell}</{tag}>")
                            html_parts.append("    </tr>")
                        html_parts.append("  </table>")
                    
                    if "image" in section:
                        html_parts.append(f"  <img src='{section['image']}' alt='Image'>")
            
            # Liste
            if "list" in data:
                html_parts.append("  <ul>")
                for item in data["list"]:
                    html_parts.append(f"    <li>{item}</li>")
                html_parts.append("  </ul>")
            
            # Pied de page
            if config.include_footer:
                footer_text = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
                html_parts.append(f"  <div class='footer'>{footer_text}</div>")
            
            html_parts.extend([
                "</body>",
                "</html>"
            ])
            
            html_content = "\n".join(html_parts)
        
        # Sauvegarder
        output.write_text(html_content, encoding="utf-8")
        
        return output
    
    async def generate_markdown_from_data(
        self,
        data: Dict[str, Any],
        output: Union[str, Path],
        config: Optional[GenerationConfig] = None
    ) -> Path:
        """Génère un fichier Markdown à partir de données"""
        output = Path(output)
        
        md_parts = []
        
        # Titre
        if "title" in data:
            md_parts.append(f"# {data['title']}\n")
        
        # Métadonnées
        if "metadata" in data:
            for key, value in data["metadata"].items():
                md_parts.append(f"**{key}:** {value}\n")
            md_parts.append("\n")
        
        # Sections
        if "sections" in data:
            for section in data["sections"]:
                if "title" in section:
                    md_parts.append(f"## {section['title']}\n")
                
                if "content" in section:
                    if isinstance(section["content"], str):
                        md_parts.append(f"{section['content']}\n")
                    elif isinstance(section["content"], list):
                        for para in section["content"]:
                            md_parts.append(f"{para}\n")
                
                if "table" in section:
                    table_data = section["table"]
                    # En-tête
                    md_parts.append("| " + " | ".join(str(cell) for cell in table_data[0]) + " |")
                    md_parts.append("| " + " | ".join("---" for _ in table_data[0]) + " |")
                    # Données
                    for row in table_data[1:]:
                        md_parts.append("| " + " | ".join(str(cell) for cell in row) + " |")
                    md_parts.append("\n")
                
                if "image" in section:
                    md_parts.append(f"![Image]({section['image']})\n")
                
                md_parts.append("\n")
        
        # Liste
        if "list" in data:
            for item in data["list"]:
                md_parts.append(f"- {item}")
            md_parts.append("\n")
        
        # Pied de page
        md_parts.append(f"\n---\n*Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}*")
        
        # Sauvegarder
        output.write_text("\n".join(md_parts), encoding="utf-8")
        
        return output
    
    async def create_report_from_dict(
        self,
        data: Dict[str, Any],
        output_format: str = "pdf",
        output: Optional[Union[str, Path]] = None,
        config: Optional[GenerationConfig] = None
    ) -> Path:
        """
        Créé un rapport complet à partir d'un dictionnaire
        
        Args:
            data: Données structurées
            output_format: Format de sortie (pdf, docx, html, xlsx, md)
            output: Chemin de sortie (optionnel)
            config: Configuration
        """
        if output is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output = Path(f"report_{timestamp}.{output_format}")
        
        output = Path(output)
        
        if output_format == "pdf":
            return await self.generate_pdf_from_data(data, output, config=config)
        elif output_format == "docx":
            return await self.generate_docx_from_data(data, output, config=config)
        elif output_format == "html":
            return await self.generate_html_from_data(data, output, config=config)
        elif output_format in ["xlsx", "excel"]:
            return await self.generate_excel_from_data(data, output, config=config)
        elif output_format in ["md", "markdown"]:
            return await self.generate_markdown_from_data(data, output, config=config)
        else:
            raise ValueError(f"Format non supporté: {output_format}")
