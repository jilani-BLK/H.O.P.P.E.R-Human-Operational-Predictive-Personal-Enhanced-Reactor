"""
HOPPER - Document Reader & Analyzer
Lecture et analyse de documents locaux et web SANS API externe
100% local et confidentiel
"""

from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass
import mimetypes
import hashlib
from datetime import datetime
from loguru import logger

# Parsers locaux
import PyPDF2
import docx
import openpyxl
import markdown
import html2text
import json
import csv
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import requests


@dataclass
class Document:
    """Document structuré"""
    id: str
    source: str  # file path ou URL
    type: str  # pdf, docx, html, etc.
    title: Optional[str] = None
    content: str = ""
    metadata: Optional[Dict[str, Any]] = None
    sections: Optional[List[Dict[str, Any]]] = None
    extracted_at: Optional[str] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        if self.sections is None:
            self.sections = []
        if self.extracted_at is None:
            self.extracted_at = datetime.now().isoformat()
        if not self.id:
            self.id = hashlib.md5(self.source.encode()).hexdigest()


class LocalDocumentReader:
    """
    Lecteur de documents 100% local
    
    Supporte:
    - PDF (PyPDF2)
    - Word (.docx)
    - Excel (.xlsx)
    - Markdown (.md)
    - HTML
    - Text (.txt, .log, .csv)
    - JSON
    - XML
    - Pages web (extraction locale)
    
    AUCUNE donnée ne quitte la machine
    """
    
    def __init__(self, llm_analyzer: Optional[Any] = None):
        """
        Args:
            llm_analyzer: Agent LLM local pour analyse sémantique (optionnel)
        """
        self.llm_analyzer = llm_analyzer
        self.supported_formats = {
            'pdf': self._read_pdf,
            'docx': self._read_docx,
            'doc': self._read_doc_fallback,
            'xlsx': self._read_xlsx,
            'xls': self._read_xls_fallback,
            'md': self._read_markdown,
            'html': self._read_html,
            'htm': self._read_html,
            'txt': self._read_text,
            'log': self._read_text,
            'csv': self._read_csv,
            'tsv': self._read_csv,
            'json': self._read_json,
            'xml': self._read_xml,
            'web': self._read_web
        }
        
        logger.info(f"DocumentReader initialisé avec {len(self.supported_formats)} formats supportés")
    
    def read(self, source: Union[str, Path], source_type: str = "auto") -> Document:
        """
        Lit un document depuis un fichier local ou URL
        
        Args:
            source: Chemin fichier ou URL
            source_type: Type (auto-détection par défaut)
        
        Returns:
            Document structuré avec contenu extrait
        """
        
        source_str = str(source)
        logger.info(f"Lecture document: {source_str}")
        
        # Détection automatique du type
        if source_type == "auto":
            if source_str.startswith(("http://", "https://")):
                source_type = "web"
            else:
                # Détection par extension
                path = Path(source_str)
                if not path.exists():
                    raise FileNotFoundError(f"Fichier introuvable: {source_str}")
                
                extension = path.suffix.lower().lstrip('.')
                source_type = extension if extension in self.supported_formats else "txt"
        
        # Sélectionner le parser approprié
        parser = self.supported_formats.get(source_type)
        if not parser:
            raise ValueError(f"Format non supporté: {source_type}")
        
        # Extraire le contenu
        doc = parser(source_str)
        
        # Analyse sémantique optionnelle avec LLM local
        if self.llm_analyzer:
            doc = self._semantic_analysis(doc)
        
        logger.success(f"Document lu: {len(doc.content)} caractères")
        return doc
    
    def _read_pdf(self, file_path: str) -> Document:
        """Extrait texte et structure d'un PDF (100% local)"""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Métadonnées
            metadata = {
                "pages": len(pdf_reader.pages),
                "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                "creator": pdf_reader.metadata.get('/Creator', '') if pdf_reader.metadata else ''
            }
            
            # Extraire texte page par page
            sections = []
            full_text = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                sections.append({
                    "type": "page",
                    "number": page_num + 1,
                    "content": text,
                    "char_count": len(text)
                })
                
                full_text.append(text)
            
            return Document(
                id="",
                source=file_path,
                type="pdf",
                title=metadata.get("title", Path(file_path).stem),
                content="\n\n".join(full_text),
                metadata=metadata,
                sections=sections
            )
    
    def _read_docx(self, file_path: str) -> Document:
        """Extrait texte et structure d'un Word .docx"""
        
        doc = docx.Document(file_path)
        
        # Métadonnées
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "sections": len(doc.sections)
        }
        
        # Extraire par paragraphes
        sections = []
        full_text = []
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                sections.append({
                    "type": "paragraph",
                    "number": i + 1,
                    "content": para.text,
                    "style": para.style.name
                })
                full_text.append(para.text)
        
        # Extraire tables
        for table_num, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            
            sections.append({
                "type": "table",
                "number": table_num + 1,
                "content": "\n".join(table_text)
            })
        
        return Document(
            id="",
            source=file_path,
            type="docx",
            title=Path(file_path).stem,
            content="\n\n".join(full_text),
            metadata=metadata,
            sections=sections
        )
    
    def _read_doc_fallback(self, file_path: str) -> Document:
        """Fallback pour .doc (ancien format)"""
        logger.warning("Format .doc ancien - conversion limitée")
        # TODO: Utiliser python-docx2txt ou antiword si disponible
        return Document(
            id="",
            source=file_path,
            type="doc",
            title=Path(file_path).stem,
            content="[Format .doc ancien - parser spécifique requis]",
            metadata={"warning": "Format non supporté nativement"}
        )
    
    def _read_xlsx(self, file_path: str) -> Document:
        """Extrait données d'un Excel .xlsx"""
        
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        metadata = {
            "sheets": len(workbook.sheetnames),
            "sheet_names": workbook.sheetnames
        }
        
        sections = []
        full_text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Extraire données
            rows_data = []
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                rows_data.append(row_text)
            
            sheet_content = "\n".join(rows_data)
            
            sections.append({
                "type": "sheet",
                "name": sheet_name,
                "rows": len(rows_data),
                "content": sheet_content
            })
            
            full_text.append(f"# {sheet_name}\n{sheet_content}")
        
        return Document(
            id="",
            source=file_path,
            type="xlsx",
            title=Path(file_path).stem,
            content="\n\n".join(full_text),
            metadata=metadata,
            sections=sections
        )
    
    def _read_xls_fallback(self, file_path: str) -> Document:
        """Fallback pour .xls (ancien format)"""
        logger.warning("Format .xls ancien - utiliser xlrd si nécessaire")
        return Document(
            id="",
            source=file_path,
            type="xls",
            title=Path(file_path).stem,
            content="[Format .xls ancien - parser spécifique requis]",
            metadata={"warning": "Format non supporté nativement"}
        )
    
    def _read_markdown(self, file_path: str) -> Document:
        """Extrait et convertit Markdown"""
        
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convertir en HTML pour extraction structure
        html_content = markdown.markdown(md_content, extensions=['extra', 'toc'])
        
        # Parser avec BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extraire sections (basé sur headers)
        sections = []
        for header in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(header.name[1])
            sections.append({
                "type": "heading",
                "level": level,
                "content": header.get_text()
            })
        
        return Document(
            id="",
            source=file_path,
            type="markdown",
            title=Path(file_path).stem,
            content=md_content,
            metadata={"sections_count": len(sections)},
            sections=sections
        )
    
    def _read_html(self, file_path: str) -> Document:
        """Extrait contenu d'un fichier HTML"""
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Métadonnées
        title = soup.title.string if soup.title else Path(file_path).stem
        
        # Extraire texte propre
        h2t = html2text.HTML2Text()
        h2t.ignore_links = False
        h2t.ignore_images = False
        text_content = h2t.handle(html_content)
        
        # Extraire structure
        sections = []
        for header in soup.find_all(['h1', 'h2', 'h3']):
            sections.append({
                "type": "heading",
                "level": int(header.name[1]),
                "content": header.get_text()
            })
        
        return Document(
            id="",
            source=file_path,
            type="html",
            title=title,
            content=text_content,
            metadata={"sections_count": len(sections)},
            sections=sections
        )
    
    def _read_text(self, file_path: str) -> Document:
        """Lit fichier texte brut"""
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Tentative de détection de structure
        lines = content.split('\n')
        sections = []
        
        # Détecter sections basées sur lignes vides
        current_section = []
        for line in lines:
            if line.strip():
                current_section.append(line)
            elif current_section:
                sections.append({
                    "type": "paragraph",
                    "content": "\n".join(current_section)
                })
                current_section = []
        
        if current_section:
            sections.append({
                "type": "paragraph",
                "content": "\n".join(current_section)
            })
        
        return Document(
            id="",
            source=file_path,
            type="text",
            title=Path(file_path).stem,
            content=content,
            metadata={"lines": len(lines), "chars": len(content)},
            sections=sections
        )
    
    def _read_csv(self, file_path: str) -> Document:
        """Lit fichier CSV/TSV"""
        
        # Détection délimiteur
        with open(file_path, 'r', encoding='utf-8') as f:
            sample = f.read(1024)
            delimiter = '\t' if '\t' in sample else ','
        
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
        
        # Première ligne = headers
        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []
        
        # Convertir en texte lisible
        content_lines = []
        for row in rows:
            content_lines.append(" | ".join(row))
        
        return Document(
            id="",
            source=file_path,
            type="csv",
            title=Path(file_path).stem,
            content="\n".join(content_lines),
            metadata={
                "rows": len(data_rows),
                "columns": len(headers),
                "headers": headers
            },
            sections=[{"type": "table", "headers": headers, "rows": data_rows}]
        )
    
    def _read_json(self, file_path: str) -> Document:
        """Lit fichier JSON"""
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convertir en texte lisible
        content = json.dumps(data, indent=2, ensure_ascii=False)
        
        return Document(
            id="",
            source=file_path,
            type="json",
            title=Path(file_path).stem,
            content=content,
            metadata={
                "type": type(data).__name__,
                "keys": list(data.keys()) if isinstance(data, dict) else None
            },
            sections=[{"type": "json", "data": data}]
        )
    
    def _read_xml(self, file_path: str) -> Document:
        """Lit fichier XML"""
        
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # Convertir en texte lisible
        def xml_to_text(element, indent=0):
            lines = []
            prefix = "  " * indent
            
            # Tag et attributs
            tag_line = f"{prefix}<{element.tag}"
            if element.attrib:
                attrs = " ".join([f'{k}="{v}"' for k, v in element.attrib.items()])
                tag_line += f" {attrs}"
            tag_line += ">"
            lines.append(tag_line)
            
            # Texte
            if element.text and element.text.strip():
                lines.append(f"{prefix}  {element.text.strip()}")
            
            # Enfants
            for child in element:
                lines.extend(xml_to_text(child, indent + 1))
            
            lines.append(f"{prefix}</{element.tag}>")
            return lines
        
        content = "\n".join(xml_to_text(root))
        
        return Document(
            id="",
            source=file_path,
            type="xml",
            title=Path(file_path).stem,
            content=content,
            metadata={"root_tag": root.tag},
            sections=[{"type": "xml", "root": root.tag}]
        )
    
    def _read_web(self, url: str) -> Document:
        """
        Extrait contenu d'une page web (100% local après téléchargement)
        
        IMPORTANT: Le HTML est téléchargé puis traité localement
        Aucune donnée sensible n'est envoyée à des API tierces
        """
        
        logger.info(f"Téléchargement page web: {url}")
        
        # Télécharger le HTML
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        html_content = response.text
        
        # Parser localement avec BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extraire titre
        title = soup.title.string if soup.title else url
        
        # Supprimer scripts et styles
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extraire texte principal
        h2t = html2text.HTML2Text()
        h2t.ignore_links = False
        h2t.body_width = 0
        text_content = h2t.handle(html_content)
        
        # Extraire structure
        sections = []
        for header in soup.find_all(['h1', 'h2', 'h3']):
            sections.append({
                "type": "heading",
                "level": int(header.name[1]),
                "content": header.get_text().strip()
            })
        
        # Extraire liens
        links = []
        for a in soup.find_all('a', href=True):
            links.append({
                "text": a.get_text().strip(),
                "url": a['href']
            })
        
        return Document(
            id="",
            source=url,
            type="web",
            title=title,
            content=text_content,
            metadata={
                "url": url,
                "sections_count": len(sections),
                "links_count": len(links)
            },
            sections=sections
        )
    
    def _semantic_analysis(self, doc: Document) -> Document:
        """
        Analyse sémantique avec LLM LOCAL
        
        Extrait:
        - Résumé
        - Thèmes principaux
        - Entités mentionnées
        - Structure logique
        
        IMPORTANT: Tout est fait localement, aucun appel API externe
        """
        
        logger.info(f"Analyse sémantique locale du document: {doc.title}")
        
        # Construire prompt pour LLM
        prompt = f"""Analyse ce document et extrais:
1. Un résumé en 2-3 phrases
2. Les 5 thèmes principaux
3. Les entités importantes (personnes, lieux, organisations, concepts)
4. La structure logique

DOCUMENT:
Titre: {doc.title}
Type: {doc.type}
Contenu (premiers 2000 caractères):
{doc.content[:2000]}

Réponds en JSON structuré.
"""
        
        # Appel au LLM local (via l'agent)
        try:
            analysis = self.llm_analyzer.analyze(prompt)
            
            # Ajouter à metadata
            doc.metadata["semantic_analysis"] = analysis
            
            logger.success("Analyse sémantique complétée")
            
        except Exception as e:
            logger.warning(f"Analyse sémantique échouée: {e}")
        
        return doc
    
    def batch_read(self, sources: List[Union[str, Path]]) -> List[Document]:
        """Lit plusieurs documents en batch"""
        
        documents = []
        for source in sources:
            try:
                doc = self.read(source)
                documents.append(doc)
            except Exception as e:
                logger.error(f"Erreur lecture {source}: {e}")
        
        logger.info(f"Batch terminé: {len(documents)}/{len(sources)} documents lus")
        return documents
    
    def export_to_json(self, doc: Document, output_path: str):
        """Exporte document analysé en JSON"""
        
        data = {
            "id": doc.id,
            "source": doc.source,
            "type": doc.type,
            "title": doc.title,
            "content": doc.content,
            "metadata": doc.metadata,
            "sections": doc.sections,
            "extracted_at": doc.extracted_at
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        logger.success(f"Document exporté: {output_path}")


# Exemple d'utilisation
if __name__ == "__main__":
    # Initialiser le reader
    reader = LocalDocumentReader()
    
    # Lire un PDF
    pdf_doc = reader.read("document.pdf")
    print(f"PDF: {pdf_doc.title} - {len(pdf_doc.content)} caractères")
    
    # Lire une page web (traitement 100% local après téléchargement)
    web_doc = reader.read("https://example.com")
    print(f"Web: {web_doc.title} - {len(web_doc.sections or [])} sections")
    
    # Batch de documents
    docs = reader.batch_read([
        "doc1.pdf",
        "doc2.docx",
        "https://example.com/page",
        "data.json"
    ])
    
    print(f"Traités: {len(docs)} documents")
